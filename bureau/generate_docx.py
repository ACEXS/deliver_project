from docx import Document
from docx.shared import Inches
import os

# 创建文档
doc = Document()

# 添加标题
doc.add_heading('泰宁县乡村振兴局后台管理系统使用说明书', 0)

# 1. 系统概述
doc.add_heading('1. 系统概述', level=1)
doc.add_paragraph('泰宁县乡村振兴局后台管理系统是一个专为泰宁县乡村振兴局设计的综合管理平台，旨在帮助管理人员高效管理农产品、订单、农户、资讯和帮扶项目等信息。系统采用现代化的Web技术开发，具有直观的用户界面和丰富的功能模块。')

# 1.1 系统功能模块
doc.add_heading('1.1 系统功能模块', level=2)
modules = [
    '**首页**：展示系统概览和关键数据',
    '**商品管理**：管理农产品信息',
    '**订单管理**：处理和跟踪订单',
    '**农户管理**：管理农户信息',
    '**资讯管理**：发布和管理资讯',
    '**数据统计**：查看系统数据统计',
    '**帮扶项目**：管理帮扶项目信息',
    '**项目投票**：支持投票选择优先实施的帮扶项目',
    '**系统设置**：配置系统参数'
]
for module in modules:
    doc.add_paragraph(module)

# 1.2 系统登录
doc.add_heading('1.2 系统登录', level=2)
doc.add_paragraph('1. 打开浏览器，访问系统网址')
doc.add_paragraph('2. 输入用户名和密码（默认：admin/admin）')
doc.add_paragraph('3. 点击登录按钮进入系统')

# 添加登录页面截图
image_path = 'image/login_page.png'
if os.path.exists(image_path):
    doc.add_picture(image_path, width=Inches(6))
doc.add_paragraph('登录页面', style='Caption')

# 2. 功能模块详细说明
doc.add_heading('2. 功能模块详细说明', level=1)

# 2.1 首页
doc.add_heading('2.1 首页', level=2)
doc.add_paragraph('首页展示系统概览，包括商品、订单、农户和销售数据的统计信息，以及最近的订单列表。')
image_path = 'image/dashboard_page.png'
if os.path.exists(image_path):
    doc.add_picture(image_path, width=Inches(6))
doc.add_paragraph('首页', style='Caption')

# 2.2 商品管理
doc.add_heading('2.2 商品管理', level=2)
doc.add_paragraph('商品管理模块用于管理农产品信息，包括添加、编辑、删除商品，以及查看商品状态和库存。')
image_path = 'image/products_page.png'
if os.path.exists(image_path):
    doc.add_picture(image_path, width=Inches(6))
doc.add_paragraph('商品管理页面', style='Caption')

# 2.3 订单管理
doc.add_heading('2.3 订单管理', level=2)
doc.add_paragraph('订单管理模块用于处理和跟踪订单，包括查看订单状态、修改订单状态等。')
image_path = 'image/orders_page.png'
if os.path.exists(image_path):
    doc.add_picture(image_path, width=Inches(6))
doc.add_paragraph('订单管理页面', style='Caption')

# 2.4 农户管理
doc.add_heading('2.4 农户管理', level=2)
doc.add_paragraph('农户管理模块用于管理农户信息，包括添加、编辑、删除农户，以及查看农户的经营范围和等级。')
image_path = 'image/farmers_page.png'
if os.path.exists(image_path):
    doc.add_picture(image_path, width=Inches(6))
doc.add_paragraph('农户管理页面', style='Caption')

# 2.5 资讯管理
doc.add_heading('2.5 资讯管理', level=2)
doc.add_paragraph('资讯管理模块用于发布和管理资讯，包括添加、编辑、删除资讯，以及查看资讯的阅读量。')
image_path = 'image/news_page.png'
if os.path.exists(image_path):
    doc.add_picture(image_path, width=Inches(6))
doc.add_paragraph('资讯管理页面', style='Caption')

# 2.6 数据统计
doc.add_heading('2.6 数据统计', level=2)
doc.add_paragraph('数据统计模块用于查看系统数据统计，包括销售趋势和分类分布。')
image_path = 'image/statistics_page.png'
if os.path.exists(image_path):
    doc.add_picture(image_path, width=Inches(6))
doc.add_paragraph('数据统计页面', style='Caption')

# 2.7 帮扶项目
doc.add_heading('2.7 帮扶项目', level=2)
doc.add_paragraph('帮扶项目模块用于管理帮扶项目信息，包括查看项目进度、实施过程和已完成事项。')
image_path = 'image/projects_page.png'
if os.path.exists(image_path):
    doc.add_picture(image_path, width=Inches(6))
doc.add_paragraph('帮扶项目页面', style='Caption')

# 2.8 项目投票
doc.add_heading('2.8 项目投票', level=2)
doc.add_paragraph('项目投票模块用于支持投票选择优先实施的帮扶项目，包括查看投票结果和图表展示。')
image_path = 'image/projectVoting_page.png'
if os.path.exists(image_path):
    doc.add_picture(image_path, width=Inches(6))
doc.add_paragraph('项目投票页面', style='Caption')

# 2.9 系统设置
doc.add_heading('2.9 系统设置', level=2)
doc.add_paragraph('系统设置模块用于配置系统参数，包括系统名称、联系方式和通知设置。')
image_path = 'image/settings_page.png'
if os.path.exists(image_path):
    doc.add_picture(image_path, width=Inches(6))
doc.add_paragraph('系统设置页面', style='Caption')

# 3. 系统操作指南
doc.add_heading('3. 系统操作指南', level=1)

# 3.1 添加商品
doc.add_heading('3.1 添加商品', level=2)
doc.add_paragraph('1. 进入商品管理页面')
doc.add_paragraph('2. 点击「新增商品」按钮')
doc.add_paragraph('3. 填写商品信息，包括商品名称、分类、价格和库存')
doc.add_paragraph('4. 点击「保存」按钮')

# 3.2 处理订单
doc.add_heading('3.2 处理订单', level=2)
doc.add_paragraph('1. 进入订单管理页面')
doc.add_paragraph('2. 查看订单列表，了解订单状态')
doc.add_paragraph('3. 点击「查看」按钮查看订单详情')
doc.add_paragraph('4. 根据需要修改订单状态')

# 3.3 管理农户
doc.add_heading('3.3 管理农户', level=2)
doc.add_paragraph('1. 进入农户管理页面')
doc.add_paragraph('2. 点击「新增农户」按钮添加农户')
doc.add_paragraph('3. 填写农户信息，包括姓名、联系电话、所在村庄和主营产品')
doc.add_paragraph('4. 点击「保存」按钮')

# 3.4 发布资讯
doc.add_heading('3.4 发布资讯', level=2)
doc.add_paragraph('1. 进入资讯管理页面')
doc.add_paragraph('2. 点击「发布资讯」按钮')
doc.add_paragraph('3. 填写资讯信息，包括标题、分类和内容')
doc.add_paragraph('4. 点击「保存」按钮')

# 3.5 管理项目
doc.add_heading('3.5 管理项目', level=2)
doc.add_paragraph('1. 进入帮扶项目页面')
doc.add_paragraph('2. 点击「新增项目」按钮添加项目')
doc.add_paragraph('3. 填写项目信息，包括项目名称、描述、预算和截止日期')
doc.add_paragraph('4. 点击「保存」按钮')

# 3.6 项目投票
doc.add_heading('3.6 项目投票', level=2)
doc.add_paragraph('1. 进入项目投票页面')
doc.add_paragraph('2. 查看项目列表和当前投票数')
doc.add_paragraph('3. 点击「投票」按钮为支持的项目投票')
doc.add_paragraph('4. 查看投票结果图表')

# 4. 系统维护
doc.add_heading('4. 系统维护', level=1)

# 4.1 密码修改
doc.add_heading('4.1 密码修改', level=2)
doc.add_paragraph('1. 进入系统设置页面')
doc.add_paragraph('2. 点击「修改密码」按钮')
doc.add_paragraph('3. 输入旧密码和新密码')
doc.add_paragraph('4. 点击「保存」按钮')

# 4.2 系统通知设置
doc.add_heading('4.2 系统通知设置', level=2)
doc.add_paragraph('1. 进入系统设置页面')
doc.add_paragraph('2. 调整通知设置，包括新订单通知、库存预警和系统公告')
doc.add_paragraph('3. 点击「保存」按钮')

# 5. 常见问题
doc.add_heading('5. 常见问题', level=1)

# 5.1 登录失败
doc.add_heading('5.1 登录失败', level=2)
doc.add_paragraph('- 检查用户名和密码是否正确')
doc.add_paragraph('- 确保网络连接正常')

# 5.2 数据不显示
doc.add_heading('5.2 数据不显示', level=2)
doc.add_paragraph('- 检查网络连接')
doc.add_paragraph('- 刷新页面')
doc.add_paragraph('- 联系系统管理员')

# 5.3 操作失败
doc.add_heading('5.3 操作失败', level=2)
doc.add_paragraph('- 确保操作步骤正确')
doc.add_paragraph('- 检查系统权限')
doc.add_paragraph('- 联系系统管理员')

# 6. 联系我们
doc.add_heading('6. 联系我们', level=1)
doc.add_paragraph('- 系统名称：泰宁县乡村振兴局后台管理系统')
doc.add_paragraph('- 联系电话：0598-1234567')
doc.add_paragraph('- 地址：福建省三明市泰宁县和平中街25号')

# 添加版本信息
doc.add_paragraph('')
doc.add_paragraph('---')
doc.add_paragraph('**版本信息**：V1.0')
doc.add_paragraph('**发布日期**：2026年3月16日')
doc.add_paragraph('**适用范围**：泰宁县乡村振兴局内部使用')

# 保存文档
doc.save('泰宁县乡村振兴局后台管理系统使用说明书.docx')
print('文档生成成功！')
