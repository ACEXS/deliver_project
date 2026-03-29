from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml.ns import qn
import os

# 获取当前目录的绝对路径
current_dir = os.path.abspath('.')

# 创建文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

# 设置中文字体
r_pr = style._element.get_or_add_rPr()
r_fonts = r_pr.get_or_add_rFonts()
r_fonts.set(qn('w:eastAsia'), '宋体')

# 设置标题
heading_style = doc.styles['Heading 1']
heading_font = heading_style.font
heading_font.name = 'Arial'
heading_font.size = Pt(16)
heading_font.bold = True

# 设置二级标题
heading2_style = doc.styles['Heading 2']
heading2_font = heading2_style.font
heading2_font.name = 'Arial'
heading2_font.size = Pt(14)
heading2_font.bold = True

# 标题行
doc.add_heading('大金湖文创特产商城产品使用说明书', 0)

# 1. 系统概述
doc.add_heading('1. 系统概述', level=1)
doc.add_paragraph('大金湖文创特产商城是一个专为福建省泰宁县大金湖旅游经济开发区管委会设计的电商平台，用于销售景区文创产品和特色特产。系统采用现代化的前端设计，具有美观、易用的用户界面。')

# 2. 功能模块
doc.add_heading('2. 功能模块', level=1)

# 2.1 首页
doc.add_heading('2.1 首页', level=2)
doc.add_paragraph('首页展示了商城的主要信息和推荐商品，包括：')
doc.add_paragraph('- 商城logo和导航栏')
doc.add_paragraph('- 英雄区展示大金湖风光')
doc.add_paragraph('- 推荐商品展示')
doc.add_paragraph('- 商城简介')

# 2.2 商品列表页
doc.add_heading('2.2 商品列表页', level=2)
doc.add_paragraph('商品列表页展示了所有商品，支持：')
doc.add_paragraph('- 商品分类筛选（全部、文创产品、特产）')
doc.add_paragraph('- 商品搜索功能')
doc.add_paragraph('- 商品卡片展示（图片、名称、价格、简介）')

# 2.3 商品详情页
doc.add_heading('2.3 商品详情页', level=2)
doc.add_paragraph('商品详情页展示了商品的详细信息，包括：')
doc.add_paragraph('- 商品图片')
doc.add_paragraph('- 商品名称和价格')
doc.add_paragraph('- 商品描述')
doc.add_paragraph('- 商品规格')
doc.add_paragraph('- 数量控制')
doc.add_paragraph('- 加入购物车功能')

# 2.4 购物车页面
doc.add_heading('2.4 购物车页面', level=2)
doc.add_paragraph('购物车页面管理用户的购物车，支持：')
doc.add_paragraph('- 商品列表展示')
doc.add_paragraph('- 数量调整')
doc.add_paragraph('- 商品删除')
doc.add_paragraph('- 总金额计算')
doc.add_paragraph('- 结算功能')

# 2.5 结算页面
doc.add_heading('2.5 结算页面', level=2)
doc.add_paragraph('结算页面用于用户填写订单信息，包括：')
doc.add_paragraph('- 收货人信息表单')
doc.add_paragraph('- 支付方式选择')
doc.add_paragraph('- 订单确认')

# 2.6 关于我们页面
doc.add_heading('2.6 关于我们页面', level=2)
doc.add_paragraph('关于我们页面展示了商城的背景和使命，包括：')
doc.add_paragraph('- 商城简介')
doc.add_paragraph('- 我们的优势')
doc.add_paragraph('- 我们的团队')

# 2.7 联系我们页面
doc.add_heading('2.7 联系我们页面', level=2)
doc.add_paragraph('联系我们页面提供了商城的联系方式，包括：')
doc.add_paragraph('- 联系信息')
doc.add_paragraph('- 留言表单')

# 2.8 登录页面
doc.add_heading('2.8 登录页面', level=2)
doc.add_paragraph('登录页面用于用户登录，包括：')
doc.add_paragraph('- 用户名输入')
doc.add_paragraph('- 密码输入')
doc.add_paragraph('- 登录按钮')
doc.add_paragraph('- 注册链接')

# 3. 使用流程
doc.add_heading('3. 使用流程', level=1)

# 3.1 浏览商品
doc.add_heading('3.1 浏览商品', level=2)
doc.add_paragraph('1. 进入首页，查看推荐商品')
doc.add_paragraph('2. 点击"商品列表"进入商品列表页')
doc.add_paragraph('3. 使用分类筛选或搜索功能找到感兴趣的商品')
doc.add_paragraph('4. 点击商品卡片进入商品详情页')

# 3.2 购买商品
doc.add_heading('3.2 购买商品', level=2)
doc.add_paragraph('1. 在商品详情页选择商品数量')
doc.add_paragraph('2. 点击"加入购物车"按钮')
doc.add_paragraph('3. 点击购物车图标进入购物车页面')
doc.add_paragraph('4. 确认商品信息和数量')
doc.add_paragraph('5. 点击"去结算"按钮进入结算页面')
doc.add_paragraph('6. 填写收货人信息')
doc.add_paragraph('7. 选择支付方式')
doc.add_paragraph('8. 点击"提交订单"按钮完成购买')

# 3.3 联系我们
doc.add_heading('3.3 联系我们', level=2)
doc.add_paragraph('1. 进入"联系我们"页面')
doc.add_paragraph('2. 填写留言表单')
doc.add_paragraph('3. 点击"提交"按钮发送留言')

# 4. 系统截图
doc.add_heading('4. 系统截图', level=1)

# 4.1 首页
doc.add_heading('4.1 首页', level=2)
try:
    image_path = os.path.join(current_dir, 'images', 'index-page-2026-03-27T16-08-32-968Z.png')
    doc.add_picture(image_path, width=Inches(6))
except Exception as e:
    doc.add_paragraph(f'图片加载失败: {e}')

# 4.2 商品列表页
doc.add_heading('4.2 商品列表页', level=2)
try:
    image_path = os.path.join(current_dir, 'images', 'products-page-2026-03-27T16-08-47-995Z.png')
    doc.add_picture(image_path, width=Inches(6))
except Exception as e:
    doc.add_paragraph(f'图片加载失败: {e}')

# 4.3 商品详情页
doc.add_heading('4.3 商品详情页', level=2)
try:
    image_path = os.path.join(current_dir, 'images', 'product-detail-page-2026-03-27T16-09-01-334Z.png')
    doc.add_picture(image_path, width=Inches(6))
except Exception as e:
    doc.add_paragraph(f'图片加载失败: {e}')

# 4.4 购物车页面
doc.add_heading('4.4 购物车页面', level=2)
try:
    image_path = os.path.join(current_dir, 'images', 'cart-page-2026-03-27T16-10-51-236Z.png')
    doc.add_picture(image_path, width=Inches(6))
except Exception as e:
    doc.add_paragraph(f'图片加载失败: {e}')

# 4.5 结算页面
doc.add_heading('4.5 结算页面', level=2)
try:
    image_path = os.path.join(current_dir, 'images', 'checkout-page-2026-03-27T16-11-01-587Z.png')
    doc.add_picture(image_path, width=Inches(6))
except Exception as e:
    doc.add_paragraph(f'图片加载失败: {e}')

# 4.6 关于我们页面
doc.add_heading('4.6 关于我们页面', level=2)
try:
    image_path = os.path.join(current_dir, 'images', 'about-page-2026-03-27T16-11-10-772Z.png')
    doc.add_picture(image_path, width=Inches(6))
except Exception as e:
    doc.add_paragraph(f'图片加载失败: {e}')

# 4.7 联系我们页面
doc.add_heading('4.7 联系我们页面', level=2)
try:
    image_path = os.path.join(current_dir, 'images', 'contact-page-2026-03-27T16-09-11-384Z.png')
    doc.add_picture(image_path, width=Inches(6))
except Exception as e:
    doc.add_paragraph(f'图片加载失败: {e}')

# 4.8 登录页面
doc.add_heading('4.8 登录页面', level=2)
try:
    image_path = os.path.join(current_dir, 'images', 'login-page-2026-03-27T16-10-41-873Z.png')
    doc.add_picture(image_path, width=Inches(6))
except Exception as e:
    doc.add_paragraph(f'图片加载失败: {e}')

# 5. 技术说明
doc.add_heading('5. 技术说明', level=1)

# 5.1 前端技术
doc.add_heading('5.1 前端技术', level=2)
doc.add_paragraph('- HTML5 + CSS3 + JavaScript')
doc.add_paragraph('- 响应式设计，支持移动端')
doc.add_paragraph('- 现代化的UI组件')
doc.add_paragraph('- 本地存储用于购物车管理')

# 5.2 后端技术
doc.add_heading('5.2 后端技术', level=2)
doc.add_paragraph('- Python + Flask')
doc.add_paragraph('- RESTful API')
doc.add_paragraph('- CORS支持')

# 6. 注意事项
doc.add_heading('6. 注意事项', level=1)
doc.add_paragraph('1. 本系统需要用户登录才能进行购物操作')
doc.add_paragraph('2. 登录状态有效期为1小时')
doc.add_paragraph('3. 商品图片可能会根据实际情况进行更新')
doc.add_paragraph('4. 如有任何问题，请联系客服人员')

# 7. 联系方式
doc.add_heading('7. 联系方式', level=1)
doc.add_paragraph('- 地址：福建省三明市泰宁县大金湖旅游经济开发区')
doc.add_paragraph('- 电话：0598-1234567')
doc.add_paragraph('- 邮箱：info@dajinlake.com')

# 保存文档
doc.save('产品使用说明书.docx')
print('产品使用说明书生成成功！')
print(f'当前目录: {current_dir}')